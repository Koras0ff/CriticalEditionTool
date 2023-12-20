import os
import csv
from collections import defaultdict
from docx import Document

# Function to load CSV data
def load_csv(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        reader = csv.reader(file)
        headers = next(reader)  # Get the headers
        data = [row for row in reader]  # Read the remaining data
    return headers, data

def create_critical_edition(csv_data, headers):
    doc = Document()
    
    for row in csv_data:
        line_variations = defaultdict(list)

        # Collect all variations of each line
        for i, line in enumerate(row):
            line_variations[line].append(headers[i])

        # Find the most common line (excluding blanks)
        most_common_line = max(line_variations.keys(), key=lambda l: (len(line_variations[l]), l), default=None)
        most_common_editions = line_variations.pop(most_common_line, [])

        # Add the most common line and its editions to the document
        if most_common_line:
            doc.add_paragraph(most_common_line)
            if most_common_editions:
                editions_str = ', '.join(most_common_editions)
                doc.add_paragraph(f"(In {editions_str}, written as: '{most_common_line}')")

        # Add other variations to the document
        for line, editions in line_variations.items():
            if line:  # If the line is not blank
                editions_str = ', '.join(editions)
                doc.add_paragraph(f"(In {editions_str}, written as: '{line}')")
            else:  # If the line is blank
                editions_str = ', '.join(editions)
                doc.add_paragraph(f"(There is no line in {editions_str})")

    return doc


# Function to process a folder of CSV files and output them as DOCX files
def process_csv_folder(source_folder_path, output_folder_path):
    # Create the output folder if it doesn't exist
    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)

    # Process each CSV file in the source folder
    for filename in os.listdir(source_folder_path):
        if filename.endswith('.csv'):
            csv_file_path = os.path.join(source_folder_path, filename)
            headers, csv_data = load_csv(csv_file_path)
            
            # Create the critical edition document
            critical_edition_doc = create_critical_edition(csv_data, headers)
            
            # Define the output DOCX file path
            output_docx_path = os.path.join(output_folder_path, f"{os.path.splitext(filename)[0]}_critical_edition.docx")
            
            # Save the DOCX file
            critical_edition_doc.save(output_docx_path)
            print(f"Processed and saved: {filename}")


# Example usage
source_folder = '/Users/enesyilandiloglu/Documents/GitHub/Kakule_1/CompLit/AlignedPoems'  # Replace with your source folder path
output_folder = '/Users/enesyilandiloglu/Documents/GitHub/Kakule_1/CompLit/CriticalEditions'  # Replace with your desired output folder path
process_csv_folder(source_folder, output_folder)
